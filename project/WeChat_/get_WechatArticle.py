from datetime import *
import wechatsogou
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
import time



def get_articles(original=True,timedel=1):
    ws_api = wechatsogou.WechatSogouAPI(captcha_break_time=3)
    accounts=['infoQ','成都美食']
    articles=[]
    for account in accounts:
        articles.extend(reformat(ws_api.get_gzh_article_by_history(account)))
        # 时间过滤，只选取规定天数以内的
        timestamp = int((datetime.now() - timedelta(days=timedel)).timestamp())
        articles = [article for article in articles if article['datetime'] > timestamp]
        #监测是否为原创
    if original:
        for article in articles:
            if article['copyright_stat'] != 100:
                articles.remove(article)
    return articles
def reformat(data):
   atcs = data.get('article')
   if atcs is not None:
       wechat_name = data.get('gzh')['wechat_name']
       for article in atcs:
           article['wechat_name'] = wechat_name
       return atcs

# 文章整合为文本
def to_msdocx(data):
    document = Document()
    header = '公众号最新文章({})'.format(datetime.now().strftime('%a, %b %d %H:%M'))
    document.add_heading(header, 0)
    for article in data:
        pubtime_stamp = article['datetime']
        timeArray = time.localtime(pubtime_stamp)
        formatTime = time.strftime("%Y-%m-%d %H:%M:%S", timeArray)
        # print(formatTime,type(formatTime))
        print(article)
        document.add_paragraph(article['title'], style='ListNumber')
        document.add_paragraph('摘要： ' + article['abstract'])
        document.add_paragraph('url： ' + article['content_url'])
        document.add_paragraph('时间： ' + str(formatTime))
        # document.add_paragraph('时间： ' + str(article['datetime']))
        document.add_paragraph('来自： ' + article['wechat_name']+'\n')
    document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.save('公众号文章{}.docx'.format(datetime.now().strftime('%m.%d')))


if __name__ == "__main__":
    articles=get_articles(timedel=3)
    to_msdocx(articles)